from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "mock-data" / "offline-knowledge-base" / "2026-q1-product-ops"


def set_document_fonts(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.size = Pt(10.5)


def add_section(document: Document, heading: str, paragraphs: list[str]) -> None:
    document.add_heading(heading, level=1)
    for text in paragraphs:
        document.add_paragraph(text)


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered_lines(document: Document, items: list[tuple[str, str]]) -> None:
    for index, (title, content) in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        title_run = paragraph.add_run(f"{index}. {title}：")
        title_run.bold = True
        paragraph.add_run(content)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = value


def build_doc_2025_quarterly() -> Document:
    document = Document()
    set_document_fonts(document)
    add_title(
        document,
        "25年产品季报",
        "用于季度工作汇报结构规划与管理层摘要写法参考的 mock 材料",
    )

    add_section(
        document,
        "一、材料用途",
        [
            "本材料用于沉淀 2025 年产品季度经营汇报的共性结构，重点支持季度工作汇报在“章节顺序、管理层摘要、亮点与风险表达”上的统一写法。",
            "在季度正式汇报场景中，建议优先参考本材料确定正文骨架，再结合当期数据和案例补足事实内容。",
        ],
    )

    add_section(
        document,
        "二、推荐结构",
        [
            "季度工作汇报建议按“管理层摘要、季度整体情况、经营数据分析、重点工作与亮点、问题与风险、下季度安排”的顺序组织。",
            "若存在跨部门协同、资源投入或客户侧进展等补充信息，可在经营分析和问题风险章节中展开，不建议单独拆出无强支撑的章节。",
        ],
    )
    add_bullet_list(
        document,
        [
            "管理层摘要：先给总体判断，再给关键数据变化和核心动作。",
            "季度整体情况：说明本季度的工作主线、关键节点和总体完成度。",
            "经营数据分析：围绕签约、回款、项目推进、投入产出展开。",
            "重点工作与亮点：突出阶段成果、标杆项目和版本发布。",
            "问题与风险：采用审慎表达，写明影响和后续动作。",
            "下季度安排：对应风险与目标，形成明确的行动闭环。",
        ],
    )

    add_section(
        document,
        "三、管理层摘要写法建议",
        [
            "管理层摘要需避免平铺罗列过程信息，应先交代总体判断，再用 2 至 3 个关键数字支撑判断，最后落到阶段性亮点和风险提醒。",
            "建议使用正式表达，例如“整体推进平稳”“经营指标保持改善”“重点项目交付风险可控”，避免使用口语化表述。",
            "摘要段长度不宜过长，一般控制在 120 至 180 字，保证管理层可以快速获得结论。",
        ],
    )

    add_section(
        document,
        "四、样本文段",
        [
            "2025 年 Q4，产品组围绕重点项目交付、版本能力完善与经营目标协同推进，整体工作按计划完成主要里程碑，核心项目交付和版本发布节奏保持稳定。",
            "从经营结果看，本季度签约与回款同步改善，重点项目总体可控，阶段性亮点体现在标杆客户突破和关键版本交付上；同时，仍需关注客户侧需求收敛和资源交叉占用带来的不确定性。",
            "下季度建议继续围绕重点项目闭环、产品能力深化、经营指标跟踪与风险项目收敛四条主线推进，保持季度经营工作的连续性与可追踪性。",
        ],
    )

    return document


def build_doc_2026_monthly() -> Document:
    document = Document()
    set_document_fonts(document)
    add_title(
        document,
        "26年产品月报",
        "用于月报向季度正式汇报汇总转换的 mock 材料",
    )

    add_section(
        document,
        "一、材料用途",
        [
            "本材料用于整理 2026 年产品月报中的固定字段和表述方式，帮助季度正式汇报在汇总月度成果时保持口径一致。",
            "当季度汇报需要回收月度材料时，可优先参考本文件中的归并原则和语言模板。",
        ],
    )

    add_section(
        document,
        "二、月报固定字段建议",
        [
            "月报建议固定呈现收入、签约、回款、项目推进、版本发布、风险事项六类内容，保证跨月可对比。",
            "对季度汇报而言，月报中的“过程性事件”可适度收敛，优先保留对经营结果有解释力的事实。",
        ],
    )
    add_bullet_list(
        document,
        [
            "经营指标：签约额、回款额、回款率、重点项目数量。",
            "业务进展：核心客户推进、标杆案例突破、重要需求上线。",
            "投入情况：研发投入、交付支持、专项资源占用。",
            "风险事项：客户侧协同、需求收敛、回款节奏、资源冲突。",
        ],
    )

    add_section(
        document,
        "三、月报汇总为季报的处理原则",
        [
            "将月报中的重复进展合并为阶段性结论，不建议按自然月顺序堆叠描述。",
            "若同一指标在月报中存在波动，应在季度汇报中突出趋势与原因，不应只写单月高低点。",
            "月报中的风险事项进入季度汇报后，应补充影响范围、趋势判断和下一步治理动作。",
        ],
    )

    add_section(
        document,
        "四、表述模板示例",
        [
            "月报表述：“3 月完成重点版本发布，新增 1 个标杆项目。”",
            "季度汇总表述：“本季度完成重点版本发布 3 次，新增标杆项目 2 个，产品能力建设与市场突破形成阶段性联动。”",
            "月报表述：“2 月客户侧需求变更较多，部分项目交付压力上升。”",
            "季度汇总表述：“季度内部分项目在客户侧需求收敛和跨团队协同上仍存在不确定性，需持续强化节奏管理与风险预警。”",
        ],
    )

    return document


def build_doc_analysis_cases() -> Document:
    document = Document()
    set_document_fonts(document)
    add_title(
        document,
        "常见经营分析案例",
        "用于季度汇报经营分析章节编排的 mock 材料",
    )

    add_section(
        document,
        "一、分析框架",
        [
            "经营分析章节建议采用“现象、原因、影响、建议动作”的闭环写法，先说明数据结果，再解释变化逻辑，最后落到业务判断。",
            "对于季度汇报，建议优先覆盖签约、回款、项目推进和资源投入四类经营要素，避免只有结果没有解释。",
        ],
    )

    add_section(
        document,
        "二、案例一：签约与回款联动分析",
        [
            "现象：季度签约额提升 18% 以上，回款率保持在 70% 以上，签约与回款节奏基本同步。",
            "原因：重点项目推进节奏更清晰，客户决策链路缩短，同时老客户增购转化效率提升。",
            "影响：签约和回款同步改善，有利于提升收入确定性和项目排期稳定性。",
            "建议动作：继续跟踪大项目签约后回款节奏，避免后续确认周期拉长。",
        ],
    )

    add_section(
        document,
        "三、案例二：项目推进情况分析",
        [
            "现象：在管重点项目 8 个，其中 6 个正常推进，2 个处于风险观察状态。",
            "原因：风险项目主要受客户侧需求变更和多方协同边界不清影响。",
            "影响：若风险项目不能及时收敛，将影响后续交付节奏与经营结果兑现。",
            "建议动作：对风险项目建立周级跟踪机制，并同步明确责任人与里程碑。",
        ],
    )

    add_section(
        document,
        "四、案例三：资源投入与阶段亮点分析",
        [
            "现象：本季度研发投入 1,420 人日，交付支持 380 人日，重点资源主要用于核心版本迭代和标杆项目交付。",
            "原因：季度内版本发布密度提升，重点客户场景需要更多研发与交付联动。",
            "影响：资源投入支撑了关键版本发布与标杆案例落地，但也带来跨项目调度压力。",
            "建议动作：下一季度需在资源规划前置、专项支持和项目优先级排序上进一步细化。",
        ],
    )

    return document


def build_doc_risk_cases() -> Document:
    document = Document()
    set_document_fonts(document)
    add_title(
        document,
        "经营问题经典case",
        "用于季度汇报风险章节与审慎表述的 mock 材料",
    )

    add_section(
        document,
        "一、风险表达基本原则",
        [
            "风险章节应采用审慎、克制、可执行的表达方式，避免使用“严重失控”“完全无法推进”等绝对化表述。",
            "建议按照“问题现象、影响范围、当前判断、后续动作”四段式组织，确保管理层可以快速判断风险性质和处置方式。",
        ],
    )

    add_section(
        document,
        "二、经典案例一：客户侧需求持续变更",
        [
            "问题现象：客户在项目实施后期持续追加新需求，影响交付边界和测试安排。",
            "影响范围：交付节奏被动拉长，研发和交付资源计划出现挤压。",
            "当前判断：属于阶段性协同问题，但若缺少统一变更机制，可能演变为结构性风险。",
            "建议表述：部分项目在客户侧需求收敛上仍存在不确定性，需通过版本边界确认和变更机制强化交付节奏管理。",
        ],
    )

    add_section(
        document,
        "三、经典案例二：回款节奏偏慢",
        [
            "问题现象：项目验收节点已完成，但客户内部付款流程推进较慢。",
            "影响范围：短期回款节奏受到影响，经营兑现滞后于项目进展。",
            "当前判断：主要属于流程性延迟，需关注是否向合同执行风险演化。",
            "建议表述：个别项目在验收后回款节奏上仍需持续跟踪，后续将重点推动关键节点闭环，降低回款滞后风险。",
        ],
    )

    add_section(
        document,
        "四、经典案例三：资源交叉占用",
        [
            "问题现象：多个重点项目和版本需求在同一时间窗口集中推进，造成核心资源交叉占用。",
            "影响范围：局部项目的交付效率和响应速度下降，影响协同稳定性。",
            "当前判断：属于资源调度类风险，需要通过优先级治理和前置规划处理。",
            "建议表述：当前部分重点任务在资源排布上仍存在交叉占用情况，需进一步优化优先级机制和专项支持安排。",
        ],
    )

    return document


def build_doc_metric_definitions() -> Document:
    document = Document()
    set_document_fonts(document)
    add_title(
        document,
        "产品经营数据汇报数据口径",
        "用于季度汇报核心指标定义与统计口径统一的 mock 材料",
    )

    add_section(
        document,
        "一、口径使用原则",
        [
            "季度经营汇报中的核心指标必须保证统计周期一致、去重规则一致、数据来源一致，避免在不同文档间直接混用。",
            "若采用签约、回款、资源投入等多类指标联合分析，需在首次出现时说明指标定义和时间范围。",
        ],
    )

    document.add_heading("二、核心指标口径表", level=1)
    add_table(
        document,
        ["指标", "定义", "统计口径", "使用提示"],
        [
            ["签约额", "季度内新签合同金额", "按合同签署日期统计，默认含税口径", "用于判断新增业务获取能力"],
            ["回款额", "季度内实际到账金额", "按财务到账日期统计", "需与签约额区分时间口径"],
            ["回款率", "回款额 / 签约额", "同一统计周期内计算", "用于观察签约兑现质量"],
            ["重点项目数", "纳入季度经营跟踪的项目数量", "按季度末在管项目口径统计", "建议拆分正常推进与风险项目"],
            ["研发投入", "季度内用于产品研发的人日总量", "按研发工时归集", "需说明是否含专项支持"],
            ["交付支持", "季度内用于实施与交付协同的人日总量", "按交付支持工时归集", "可与重点项目推进情况联动分析"],
        ],
    )

    add_section(
        document,
        "三、图表与引用建议",
        [
            "在季度汇报中，签约额与回款额建议成组呈现，用于体现经营兑现节奏。",
            "项目推进情况建议使用“总项目数 + 正常推进数 + 风险项目数”的方式写入正文或图表说明。",
            "资源投入类指标应与阶段性成果一起解释，避免形成“投入很多但结论不清”的展示效果。",
        ],
    )

    add_section(
        document,
        "四、推荐引用表达",
        [
            "本季度累计签约额按合同签署日期统计，累计回款额按财务到账日期统计，整体回款率在同一季度口径下计算。",
            "研发投入和交付支持按季度内实际工时归集，用于反映重点项目和版本建设的资源投入强度。",
        ],
    )

    return document


DOC_BUILDERS = [
    ("25年产品季报.docx", build_doc_2025_quarterly),
    ("26年产品月报.docx", build_doc_2026_monthly),
    ("常见经营分析案例.docx", build_doc_analysis_cases),
    ("经营问题经典case.docx", build_doc_risk_cases),
    ("产品经营数据汇报数据口径.docx", build_doc_metric_definitions),
]


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    for filename, builder in DOC_BUILDERS:
        document = builder()
        document.save(OUTPUT_DIR / filename)

    print(f"Generated {len(DOC_BUILDERS)} documents in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
